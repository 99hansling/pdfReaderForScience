import re
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog

def extract_all_citations(doc):
    """
    从文档中提取所有【】之间的内容
    """
    # 收集所有文本，保持原始格式
    full_text = ''
    
    # 从段落中收集
    for para in doc.paragraphs:
        full_text += para.text + '\n'
    
    # 从表格中收集
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + '\n'
    
    print("\n原始文本：")
    print("-" * 50)
    print(full_text)
    print("-" * 50)
    
    # 提取【】中的引用，保持原始格式
    pattern = r'【([\s\S]*?)】'
    citations = []
    
    for match in re.finditer(pattern, full_text):
        # 保持原始格式，只清理首尾空白
        citation = match.group(1).strip()
        if citation:
            citations.append(citation)
            print(f"\n找到引用：\n{citation}\n{'='*50}")
    
    return citations, full_text

def replace_citations_in_text(text, citation_map):
    """
    替换文本中的【】引用，处理多行情况
    """
    def normalize_text(t):
        """标准化文本以进行比较"""
        return ' '.join(line.strip() for line in t.split('\n') if line.strip())
    
    # 找到所有【】引用
    pattern = r'【[\s\S]*?】'
    matches = list(re.finditer(pattern, text))
    
    # 从后向前替换，避免位置变化
    for match in reversed(matches):
        full_citation = match.group(0)  # 完整的引用文本，包括【】
        citation_text = full_citation[1:-1]  # 去除【】
        normalized_citation = normalize_text(citation_text)
        
        # 查找对应的编号
        for citation, number in citation_map.items():
            if normalize_text(citation) == normalized_citation:
                # 替换整个匹配块，包括换行符
                start = match.start()
                end = match.end()
                text = text[:start] + f"[{number}]" + text[end:]
                print(f"替换【】引用：\n原文：{full_citation}\n替换为：[{number}]")
                break
    
    return text

def process_word_file():
    """
    主函数：处理Word文件中的引用
    """
    # 选择Word文件
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        title="选择Word文件",
        filetypes=[("Word文件", "*.docx")]
    )
    
    if not file_path:
        print("未选择文件")
        return
    
    try:
        # 读取Word文档
        print("正在读取Word文件...")
        doc = Document(file_path)
        
        # 提取所有引用和原始文本
        print("\n提取引用...")
        citations, full_text = extract_all_citations(doc)
        
        # 创建引用映射
        citation_map = {citation: i+1 for i, citation in enumerate(citations)}
        
        if not citation_map:
            print("未找到任何引用")
            return
            
        print(f"\n找到 {len(citation_map)} 个引用")
        
        # 替换引用
        processed_text = replace_citations_in_text(full_text, citation_map)
        
        # 创建新文档
        new_doc = Document()
        
        # 按行分割处理后的文本并写入新文档
        for line in processed_text.split('\n'):
            if line.strip():
                para = new_doc.add_paragraph(line)
                # 设置标题样式
                if line.strip().lower() in ['abstract', 'introduction', 'conclusions']:
                    para.style = 'Heading 2'
                elif line.strip().lower() == 'references':
                    para.style = 'Heading 1'
        
        # 添加参考文献部分
        new_doc.add_paragraph()
        ref_heading = new_doc.add_paragraph("References")
        ref_heading.style = 'Heading 1'
        new_doc.add_paragraph("=" * 50)
        
        # 按序号添加参考文献
        for citation, number in sorted(citation_map.items(), key=lambda x: x[1]):
            new_doc.add_paragraph(f"[{number}] {citation}")
        
        # 保存新文档
        output_path = os.path.join(os.path.dirname(file_path), 
                                 f"{os.path.splitext(os.path.basename(file_path))[0]}_processed.docx")
        new_doc.save(output_path)
        
        print(f"\n处理完成，新文件已保存到：{output_path}")
        print("\n引用映射关系：")
        for citation, number in sorted(citation_map.items(), key=lambda x: x[1]):
            print(f"[{number}] -> {citation}")
        
    except Exception as e:
        print(f"处理文件时出错: {str(e)}")

if __name__ == "__main__":
    process_word_file() 