import os
import PyPDF2
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog
from rename_by_title import extract_title
from batchGetSection import extract_sections

def create_word_document(folder_name):
    """
    创建Word文档并设置基本格式
    :param folder_name: 文件夹名称（用作文档标题）
    :return: Document对象
    """
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(folder_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加空行
    doc.add_paragraph()
    return doc

def read_pdf(file_path):
    """
    读取PDF文件并提取文本
    :param file_path: PDF文件路径
    :return: 提取的文本内容
    """
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        all_text = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text.strip():
                all_text.append(text.strip())
        return '\n'.join(all_text)

def add_paper_content(doc, title, sections):
    """
    将论文内容添加到Word文档
    :param doc: Document对象
    :param title: 论文标题
    :param sections: 论文各部分内容
    """
    # 添加标题
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加各个部分
    sections_order = ['abstract', 'introduction', 'conclusions']
    section_names = {
        'abstract': 'Abstract',
        'introduction': 'Introduction',
        'conclusions': 'Conclusions'
    }
    
    for section in sections_order:
        if sections.get(section):
            # 添加章节标题
            doc.add_heading(section_names[section], level=2)
            # 添加章节内容
            doc.add_paragraph(sections[section])
            # 添加空行
            doc.add_paragraph()
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    doc.add_paragraph()

def process_pdfs():
    """
    主函数：处理多个PDF文件并生成Word文档
    """
    # 选择PDF文件
    root = tk.Tk()
    root.withdraw()
    file_paths = filedialog.askopenfilenames(
        title="选择PDF文件",
        filetypes=[("PDF文件", "*.pdf")]
    )
    
    if not file_paths:
        print("未选择文件")
        return
    
    # 获取PDF所在的文件夹路径和名称
    folder_path = os.path.dirname(os.path.abspath(file_paths[0]))
    folder_name = os.path.basename(folder_path)
    
    print(f"\n处理文件夹: {folder_path}")
    print(f"文件夹名称: {folder_name}")
    
    # 创建Word文档
    doc = create_word_document(folder_name)
    
    # 处理每个PDF文件
    for file_path in file_paths:
        try:
            # 确保文件在同一个文件夹
            if os.path.dirname(os.path.abspath(file_path)) != folder_path:
                print(f"警告：跳过不在同一文件夹的文件: {file_path}")
                continue
                
            print(f"\n处理文件: {os.path.basename(file_path)}")
            
            # 读取PDF
            text = read_pdf(file_path)
            
            # 提取标题
            title = extract_title(text)
            if not title:
                print(f"无法提取标题，跳过文件: {file_path}")
                continue
            
            # 提取各个部分
            sections = extract_sections(text)
            
            # 添加到Word文档
            add_paper_content(doc, title, sections)
            
            print(f"已处理: {title}")
            
        except Exception as e:
            print(f"处理文件时出错 {file_path}: {str(e)}")
    
    # 保存Word文档到PDF所在的文件夹
    output_path = os.path.join(folder_path, f"{folder_name}.docx")
    doc.save(output_path)
    print(f"\nWord文档已保存: {output_path}")

if __name__ == "__main__":
    process_pdfs() 