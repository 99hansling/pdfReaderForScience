import PyPDF2
from tkinter import filedialog
import tkinter as tk
from docx import Document
import os

def clean_text(text):
    """
    清理文本，去除多余的空格、换行和非法字符
    :param text: 原始文本
    :return: 清理后的文本
    """
    # 移除控制字符（除了换行和制表符）
    text = ''.join(char for char in text if char >= ' ' or char in '\n\t')
    
    # 将多个空格替换为单个空格
    text = ' '.join(text.split())
    
    # 去除段落之间多余的换行
    text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
    
    # 确保文本是有效的XML字符
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
    
    return text

def format_research_paper(text):
    """
    将研究文章按照标准格式进行分段处理
    :param text: 清理后的原始文本
    :return: 格式化后的文本
    """
    # 分割文本为段落
    paragraphs = text.split('\n')
    if not paragraphs:
        return ""

    # 定义章节关键词
    sections = {
        '摘要': ['摘要', 'abstract'],
        '关键词': ['关键词', '关键字', 'keywords'],
        '引言': ['引言', '前言', 'introduction'],
        '研究方法': ['研究方法', '方法', '材料与方法', 'methods', 'materials and methods'],
        '结果': ['研究结果', '结果', 'results'],
        '讨论': ['讨论', '分析', 'discussion'],
        '结论': ['结论', '总结', 'conclusion'],
        '参考文献': ['参考文献', 'references']
    }

    formatted_text = []
    current_section = None
    current_text = []
    is_title = True  # 标记是否在标题部分

    # 处理每个段落
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # 检查是否是章节标题
        is_section = False
        for section, keywords in sections.items():
            if any(keyword in para.lower() for keyword in keywords):
                # 如果是第一个找到的章节（通常是摘要），先处理之前的内容作为标题
                if is_title:
                    if current_text:
                        formatted_text.append("标题\n====")
                        formatted_text.append('\n'.join(current_text))
                    is_title = False
                    current_text = []

                # 保存之前章节的内容
                if current_section and current_text:
                    formatted_text.append(f"\n{current_section}\n{'='*len(current_section)}")
                    formatted_text.append('\n'.join(current_text))

                current_section = section
                current_text = []
                is_section = True
                print(f"找到章节: {section}")
                break

        # 如果不是章节标题，添加到当前文本
        if not is_section:
            current_text.append(para)

    # 处理最后一个章节的内容
    if is_title and current_text:
        # 如果整篇文章都没有找到章节，把所有内容作为标题
        formatted_text.append("标题\n====")
        formatted_text.append('\n'.join(current_text))
    elif current_section and current_text:
        # 保存最后一个章节的内容
        formatted_text.append(f"\n{current_section}\n{'='*len(current_section)}")
        formatted_text.append('\n'.join(current_text))

    # 合并所有文本
    final_text = '\n\n'.join(formatted_text)
    print(f"\n最终文本长度: {len(final_text)}")

    # 如果没有识别到任何内容，返回原始文本
    if len(final_text.strip()) == 0:
        print("未识别到章节，使用原始文本")
        final_text = text

    return final_text

def read_pdf(pdf_path):
    """
    读取PDF文件并提取文本
    :param pdf_path: PDF文件路径
    :return: 提取的文本内容
    """
    with open(pdf_path, 'rb') as file:
        # 创建PDF读取器对象
        pdf_reader = PyPDF2.PdfReader(file)
        
        # 获取PDF页数
        num_pages = len(pdf_reader.pages)
        
        # 存储所有文本
        all_text = []
        
        # 遍历每一页并提取文本
        for page_num in range(num_pages):
            # 获取当前页面
            page = pdf_reader.pages[page_num]
            # 提取文本并添加到结果中
            page_text = page.extract_text()
            if page_text.strip():
                all_text.append(page_text.strip())
        
        # 合并所有页面的文本
        text = '\n'.join(all_text)
        
        # 打印原始文本的前200个字符（用于调试）
        print("\n提取的原始文本（前200个字符）：")
        print(text[:200])
        print("-" * 50)
        
        # 清理文本
        text = clean_text(text)
        print("\n清理后的文本（前200个字符）：")
        print(text[:200])
        print("-" * 50)
        
        # 格式化研究文章
        formatted_text = format_research_paper(text)
        return formatted_text

def select_pdf_file():
    """
    打开文件选择对话框选择PDF文件
    :return: 选择的文件路径
    """
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口
    file_path = filedialog.askopenfilename(
        title="选择PDF文件",
        filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")]
    )
    return file_path

def save_to_word(text, pdf_path):
    """
    将文本保存到Word文件
    :param text: 要保存的文本内容
    :param pdf_path: 原PDF文件路径（用于生成Word文件名）
    """
    # 创建一个新的Document对象
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = 140000  # 相当于12号字
    
    # 分割格式化后的文本
    sections = text.split('\n\n')
    
    # 处理第一个部分（标题）
    if sections and '标题' in sections[0]:
        title_text = sections[0].split('\n')[-1].strip()  # 获取标题文本
        # 添加标题
        title_para = doc.add_paragraph()
        title_para.alignment = 1  # 居中对齐
        title_run = title_para.add_run(title_text)
        title_run.font.name = '黑体'
        title_run.font.size = 168000  # 14号字
        title_run.font.bold = True
        # 添加标题后的空行
        doc.add_paragraph()
        sections = sections[1:]  # 移除已处理的标题部分
    
    # 处理其余部分
    for section in sections:
        if '=' in section:  # 这是一个章节标题
            # 添加空行
            doc.add_paragraph()
            # 处理标题
            title = section.split('\n')[0].strip()
            heading = doc.add_heading(level=1)
            heading.paragraph_format.space_before = 0
            heading.paragraph_format.space_after = 140000  # 添加标题后的间距
            run = heading.add_run(title)
            run.font.name = '黑体'
            run.font.size = 168000  # 相当于14号字
            run.font.bold = True
            # 添加标题后的空行
            doc.add_paragraph()
        else:
            # 处理正文
            paragraphs = section.split('\n')
            for p in paragraphs:
                if p.strip():
                    para = doc.add_paragraph()
                    para.paragraph_format.first_line_indent = 420000  # 首行缩进2个字符
                    para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
                    para.paragraph_format.space_after = 0  # 段后间距
                    run = para.add_run(p.strip())
                    run.font.name = '宋体'
    
    # 获取PDF文件名（不含路径）
    pdf_filename = os.path.basename(pdf_path)
    current_dir = os.path.dirname(os.path.abspath(__file__))
    word_filename = os.path.splitext(pdf_filename)[0] + '.docx'
    word_path = os.path.join(current_dir, word_filename)
    
    # 保存文档
    doc.save(word_path)
    return word_path

def main():
    print("请选择要读取的PDF文件...")
    pdf_path = select_pdf_file()
    
    if pdf_path:  # 如果用户选择了文件
        try:
            extracted_text = read_pdf(pdf_path)
            print("\n提取的文本内容：")
            print("-" * 50)  # 添加分隔线
            print(extracted_text[:200])
            print("-" * 50)  # 添加分隔线
            
            # 保存到Word文件
            word_path = save_to_word(extracted_text, pdf_path)
            print(f"\n文本已保存到Word文件：{word_path}")
            
        except Exception as e:
            print(f"发生错误：{str(e)}")
    else:
        print("未选择文件")

if __name__ == "__main__":
    main()
